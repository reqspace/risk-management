"""
Flask server for the Risk Management System.
Provides REST API endpoints to process content and extract risks/tasks.
Includes folder watcher for automatic transcript processing.
"""

import os
import sys
import threading
import time
import atexit
from pathlib import Path
from flask import Flask, request, jsonify
from dotenv import load_dotenv
from watchdog.observers.polling import PollingObserver as Observer
from watchdog.events import FileSystemEventHandler
from apscheduler.schedulers.background import BackgroundScheduler

from process import process_content
from daily_digest import send_digest_email, preview_digest


# File reading utilities
def read_text_file(file_path):
    """Read contents of a .txt or .md file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def read_docx_file(file_path):
    """Read contents of a .docx file."""
    from docx import Document
    doc = Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])


def read_file_contents(file_path):
    """Read file contents based on extension."""
    ext = Path(file_path).suffix.lower()
    if ext in ['.txt', '.md']:
        return read_text_file(file_path)
    elif ext == '.docx':
        return read_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


class TranscriptHandler(FileSystemEventHandler):
    """Watches for new transcript files and processes them."""

    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.docx'}

    def __init__(self, project_code):
        self.project_code = project_code
        self.processed_files = set()

    def on_created(self, event):
        """Handle new file creation."""
        if event.is_directory:
            return

        file_path = Path(event.src_path)

        # Check if file type is supported
        if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            return

        # Avoid processing the same file twice
        if str(file_path) in self.processed_files:
            return

        # Wait for file to be fully written
        time.sleep(1)

        self.process_transcript(file_path)

    def process_transcript(self, file_path):
        """Process a transcript file."""
        print(f"\n[Folder Watcher] New file detected: {file_path.name}", flush=True)

        try:
            # Read file contents
            content = read_file_contents(str(file_path))

            if not content.strip():
                print(f"[Folder Watcher] Skipping empty file: {file_path.name}", flush=True)
                return

            print(f"[Folder Watcher] Processing {file_path.name}...", flush=True)

            # Process through the same logic as the API
            result = process_content(
                project_code=self.project_code,
                content=content,
                source_type="meeting",
                source_name=file_path.name
            )

            self.processed_files.add(str(file_path))

            if result['success']:
                print(f"[Folder Watcher] Successfully processed {file_path.name}", flush=True)
                if result.get('changes'):
                    for change in result['changes']:
                        print(f"  - {change}", flush=True)
                summary = result.get('summary', {})
                print(f"  Summary: {summary.get('risks_added', 0)} risks, "
                      f"{summary.get('tasks_added', 0)} tasks, "
                      f"{summary.get('decisions_found', 0)} decisions", flush=True)
            else:
                print(f"[Folder Watcher] Error processing {file_path.name}: {result.get('error')}", flush=True)

        except Exception as e:
            print(f"[Folder Watcher] Error reading {file_path.name}: {e}", flush=True)


def start_folder_watcher(watch_path, project_code):
    """Start the folder watcher in a background thread."""
    watch_path = Path(watch_path)

    if not watch_path.exists():
        watch_path.mkdir(parents=True, exist_ok=True)
        print(f"[Folder Watcher] Created directory: {watch_path}")

    event_handler = TranscriptHandler(project_code)
    observer = Observer()
    observer.schedule(event_handler, str(watch_path), recursive=False)
    observer.start()

    print(f"[Folder Watcher] Watching {watch_path} for new transcripts...")
    print(f"[Folder Watcher] Supported files: .txt, .md, .docx")

    return observer

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Enable CORS for React dashboard and Electron app
from flask_cors import CORS
CORS(app, origins=['http://localhost:3000', 'http://localhost:5173', 'file://', 'null'])


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint."""
    return jsonify({"status": "healthy", "service": "risk-management-system"})


def extract_text_from_attachment(file_path: Path, extension: str) -> str:
    """Extract text content from attachment based on file type."""
    try:
        if extension in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif extension == '.docx':
            from docx import Document
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])

        elif extension == '.doc':
            try:
                import textract
                return textract.process(str(file_path)).decode('utf-8')
            except:
                return f"[Could not extract .doc file]"

        elif extension == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = []
                    for page in reader.pages:
                        text.append(page.extract_text() or '')
                    return '\n'.join(text)
            except:
                return f"[Could not extract PDF]"

        return ""
    except Exception as e:
        return f"[Error extracting {extension}: {e}]"


def process_attachments(attachments: list) -> tuple:
    """
    Process attachments from Power Automate payload.
    Returns (extracted_text, attachment_notes) tuple.
    """
    import base64
    import tempfile

    extracted_texts = []
    attachment_notes = []
    supported_extensions = ['.txt', '.md', '.docx', '.pdf', '.doc']

    for att in attachments:
        filename = att.get('filename', att.get('name', 'unknown'))
        content_b64 = att.get('content', att.get('contentBytes', ''))
        content_type = att.get('content_type', att.get('contentType', ''))

        extension = Path(filename).suffix.lower()

        # Skip Excel files - note them but don't process
        if extension in ['.xlsx', '.xls']:
            attachment_notes.append(f"[Attachment: {filename} - Excel file noted but not processed]")
            continue

        # Skip unsupported types
        if extension not in supported_extensions:
            attachment_notes.append(f"[Attachment: {filename} - Unsupported format {extension}]")
            continue

        try:
            # Decode base64 content
            content_bytes = base64.b64decode(content_b64)

            # Write to temp file for processing
            with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp:
                tmp.write(content_bytes)
                tmp_path = Path(tmp.name)

            # Extract text
            text = extract_text_from_attachment(tmp_path, extension)

            # Clean up temp file
            tmp_path.unlink()

            if text and len(text.strip()) > 10:
                extracted_texts.append(f"\n\n--- ATTACHMENT: {filename} ---\n{text}")
                attachment_notes.append(f"[Processed attachment: {filename} ({len(text)} chars)]")
                print(f"[Process] Extracted {len(text)} chars from {filename}")
            else:
                attachment_notes.append(f"[Attachment: {filename} - No text content extracted]")

        except Exception as e:
            attachment_notes.append(f"[Attachment: {filename} - Error: {str(e)}]")
            print(f"[Process] Error processing attachment {filename}: {e}")

    return '\n'.join(extracted_texts), attachment_notes


@app.route('/process', methods=['POST'])
def process_endpoint():
    """
    Process content to extract risks, tasks, and decisions.

    Expected JSON payload:
    {
        "project": "HB",
        "content": "Meeting transcript or email content...",
        "source_type": "email" or "meeting",
        "source_name": "Weekly Standup 2025-01-01",
        "attachments": [  // OPTIONAL - for email attachments
            {
                "filename": "notes.docx",
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": "base64encodedcontent..."
            }
        ]
    }

    Returns:
    {
        "success": true/false,
        "error": null or error message,
        "changes": ["Added Risk R-2025-001: ...", ...],
        "summary": {
            "risks_added": 2,
            "tasks_added": 3,
            "decisions_found": 1
        },
        "attachments_processed": ["notes.docx (1234 chars)"]
    }
    """
    # Validate request
    if not request.is_json:
        return jsonify({
            "success": False,
            "error": "Request must be JSON",
            "changes": []
        }), 400

    data = request.get_json()

    # Validate required fields
    required_fields = ['project', 'content', 'source_type', 'source_name']
    missing_fields = [f for f in required_fields if f not in data or not data[f]]

    if missing_fields:
        return jsonify({
            "success": False,
            "error": f"Missing required fields: {', '.join(missing_fields)}",
            "changes": []
        }), 400

    # Validate source_type
    valid_source_types = ['email', 'meeting', 'document', 'chat', 'other', 'email_attachment']
    if data['source_type'] not in valid_source_types:
        return jsonify({
            "success": False,
            "error": f"Invalid source_type. Must be one of: {', '.join(valid_source_types)}",
            "changes": []
        }), 400

    # Process attachments if provided
    combined_content = data['content']
    attachment_notes = []

    if data.get('attachments'):
        print(f"[Process] Processing {len(data['attachments'])} attachments...")
        attachment_text, attachment_notes = process_attachments(data['attachments'])
        if attachment_text:
            combined_content = data['content'] + attachment_text

    # Process the content
    result = process_content(
        project_code=data['project'],
        content=combined_content,
        source_type=data['source_type'],
        source_name=data['source_name']
    )

    # Return appropriate status code
    status_code = 200 if result['success'] else 500

    # Remove extracted_data from response to keep it concise (it's logged in Excel)
    response = {
        "success": result['success'],
        "error": result['error'],
        "changes": result['changes'],
        "summary": result.get('summary', {}),
        "attachments_processed": attachment_notes
    }

    return jsonify(response), status_code


@app.route('/projects', methods=['GET'])
def list_projects():
    """List available projects based on existing Risk Register files."""
    from pathlib import Path

    base_path = Path(__file__).parent.resolve()
    risk_registers_path = base_path / "Risk_Registers"

    projects = []
    if risk_registers_path.exists():
        for file in risk_registers_path.glob("Risk_Register_*.xlsx"):
            # Extract project code from filename
            project_code = file.stem.replace("Risk_Register_", "")
            projects.append({
                "code": project_code,
                "risk_register": str(file)
            })

    return jsonify({
        "projects": projects,
        "count": len(projects)
    })


@app.route('/digest', methods=['POST'])
def send_digest_endpoint():
    """Manually trigger sending the daily digest email."""
    project = request.args.get('project', 'HB')

    try:
        success = send_digest_email(project)
        if success:
            return jsonify({
                "success": True,
                "message": f"Digest email sent for project {project}"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Failed to send email. Check EMAIL_PASSWORD in .env"
            }), 500
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/digest/preview', methods=['GET'])
def preview_digest_endpoint():
    """Generate and return a preview of the digest email."""
    project = request.args.get('project', 'HB')

    try:
        from daily_digest import generate_html_email
        html = generate_html_email(project)
        return html, 200, {'Content-Type': 'text/html'}
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


# ============== Dashboard API Endpoints ==============

def read_excel_data(project_code='HB'):
    """Read all data from the Risk Register Excel file."""
    import openpyxl
    from datetime import datetime

    base_path = Path(__file__).parent.resolve()
    file_path = base_path / 'Risk_Registers' / f'Risk_Register_{project_code}.xlsx'

    if not file_path.exists():
        return {'risks': [], 'tasks': [], 'updates': []}

    wb = openpyxl.load_workbook(file_path)

    def parse_date(val):
        if val is None:
            return None
        if isinstance(val, datetime):
            return val.strftime('%Y-%m-%d')
        return str(val)

    # Read risks
    risks = []
    risk_sheet = wb['Risk Register']
    headers = [cell.value for cell in risk_sheet[1]]

    for row_num in range(3, risk_sheet.max_row + 1):
        row_data = {}
        for i, cell in enumerate(risk_sheet[row_num]):
            if i < len(headers) and headers[i]:
                val = cell.value
                if headers[i] in ['Date Identified', 'Last Updated', 'Closed Date']:
                    val = parse_date(val)
                row_data[headers[i]] = val
        if row_data.get('Risk ID'):
            risks.append(row_data)

    # Read tasks
    tasks = []
    task_sheet = wb['Tasks']
    task_headers = [cell.value for cell in task_sheet[1]]

    for row_num in range(3, task_sheet.max_row + 1):
        row_data = {}
        for i, cell in enumerate(task_sheet[row_num]):
            if i < len(task_headers) and task_headers[i]:
                val = cell.value
                if task_headers[i] in ['Due Date', 'Created Date', 'Completed Date']:
                    val = parse_date(val)
                row_data[task_headers[i]] = val
        if row_data.get('Task ID'):
            tasks.append(row_data)

    # Read update log
    updates = []
    update_sheet = wb['Update Log']
    update_headers = [cell.value for cell in update_sheet[1]]

    for row_num in range(3, update_sheet.max_row + 1):
        row_data = {}
        for i, cell in enumerate(update_sheet[row_num]):
            if i < len(update_headers) and update_headers[i]:
                row_data[update_headers[i]] = cell.value
        if row_data.get('Timestamp'):
            updates.append(row_data)

    # Read milestones (P6 baseline data)
    milestones = []
    if 'Milestones' in wb.sheetnames:
        ms_sheet = wb['Milestones']
        ms_headers = [cell.value for cell in ms_sheet[1]]

        for row_num in range(2, ms_sheet.max_row + 1):
            row_data = {}
            for i, cell in enumerate(ms_sheet[row_num]):
                if i < len(ms_headers) and ms_headers[i]:
                    val = cell.value
                    if ms_headers[i] in ['Baseline Date', 'Current Date']:
                        val = parse_date(val)
                    row_data[ms_headers[i]] = val
            if row_data.get('Milestone ID') or row_data.get('Milestone'):
                milestones.append(row_data)

    wb.close()
    return {'risks': risks, 'tasks': tasks, 'updates': updates, 'milestones': milestones}


def get_all_projects():
    """Get list of all project codes from Risk Register files."""
    base_path = Path(__file__).parent.resolve()
    risk_registers_path = base_path / "Risk_Registers"
    projects = []
    if risk_registers_path.exists():
        for file in risk_registers_path.glob("Risk_Register_*.xlsx"):
            project_code = file.stem.replace("Risk_Register_", "")
            projects.append(project_code)
    return sorted(projects)


def get_project_stats(project_code):
    """Get stats for a single project."""
    from datetime import datetime
    data = read_excel_data(project_code)
    risks = data['risks']
    tasks = data['tasks']
    today = datetime.now().date()

    active_risks = len([r for r in risks if r.get('Status') in ['Open', 'Active', 'Escalated']])
    watching_risks = len([r for r in risks if r.get('Status') == 'Watching'])
    closed_risks = len([r for r in risks if r.get('Status') == 'Closed'])
    open_tasks = [t for t in tasks if t.get('Status') not in ['Completed', 'Done']]

    overdue_count = 0
    for t in open_tasks:
        due_date = t.get('Due Date')
        if due_date:
            try:
                if isinstance(due_date, str):
                    due_date = datetime.strptime(due_date, '%Y-%m-%d').date()
                elif hasattr(due_date, 'date'):
                    due_date = due_date.date()
                if due_date < today:
                    overdue_count += 1
            except:
                pass

    high_priority = len([r for r in risks if r.get('Probability') == 'High' and r.get('Status') != 'Closed'])

    # Calculate health status
    if high_priority >= 3 or overdue_count >= 3:
        health = 'Critical'
    elif high_priority >= 1 or active_risks >= 5:
        health = 'At Risk'
    elif active_risks > 0:
        health = 'Caution'
    else:
        health = 'Healthy'

    return {
        'active_risks': active_risks,
        'watching_risks': watching_risks,
        'closed_risks': closed_risks,
        'total_risks': len(risks),
        'open_tasks': len(open_tasks),
        'overdue_tasks': overdue_count,
        'total_tasks': len(tasks),
        'high_priority': high_priority,
        'health': health
    }


@app.route('/api/portfolio', methods=['GET'])
def api_get_portfolio():
    """Get all projects with their stats for portfolio view."""
    try:
        projects = get_all_projects()
        portfolio = []

        for project_code in projects:
            try:
                stats = get_project_stats(project_code)
                portfolio.append({
                    'code': project_code,
                    'stats': stats
                })
            except Exception as e:
                print(f"Error loading project {project_code}: {e}")

        return jsonify({
            "success": True,
            "projects": portfolio,
            "count": len(portfolio)
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/risks', methods=['GET'])
def api_get_risks():
    """Get risks from the Risk Register. If no project specified, returns all."""
    project = request.args.get('project')
    try:
        if project:
            data = read_excel_data(project)
            risks = data['risks']
            for r in risks:
                r['project'] = project
        else:
            # Get all risks from all projects
            risks = []
            for proj in get_all_projects():
                data = read_excel_data(proj)
                for r in data['risks']:
                    r['project'] = proj
                    risks.append(r)

        return jsonify({
            "success": True,
            "risks": risks,
            "count": len(risks)
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/tasks', methods=['GET'])
def api_get_tasks():
    """Get tasks from the Risk Register. If no project specified, returns all."""
    project = request.args.get('project')
    try:
        if project:
            data = read_excel_data(project)
            tasks = data['tasks']
            for t in tasks:
                t['project'] = project
        else:
            tasks = []
            for proj in get_all_projects():
                data = read_excel_data(proj)
                for t in data['tasks']:
                    t['project'] = proj
                    tasks.append(t)

        return jsonify({
            "success": True,
            "tasks": tasks,
            "count": len(tasks)
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/tasks', methods=['POST'])
def api_create_task():
    """Create a new task in the Risk Register and optionally in Outlook."""
    try:
        data = request.get_json()
        project = data.get('project', 'HB')
        task_text = data.get('task')
        owner = data.get('owner')
        due_date = data.get('due_date')
        add_to_outlook = data.get('add_to_outlook', False)

        if not task_text:
            return jsonify({"success": False, "error": "Task description required"}), 400

        # Load the workbook
        file_path = get_risk_register_path(project)
        if not file_path.exists():
            return jsonify({"success": False, "error": f"Project {project} not found"}), 404

        import openpyxl
        wb = openpyxl.load_workbook(file_path)
        task_sheet = wb['Tasks']

        # Find the next task ID
        existing_ids = []
        for row in range(3, task_sheet.max_row + 1):
            task_id = task_sheet.cell(row=row, column=1).value
            if task_id and task_id.startswith(f'{project}-T'):
                try:
                    num = int(task_id.split('-T')[1])
                    existing_ids.append(num)
                except:
                    pass

        next_num = max(existing_ids) + 1 if existing_ids else 1
        new_task_id = f"{project}-T{next_num:03d}"

        # Find the next empty row
        next_row = task_sheet.max_row + 1
        for row in range(3, task_sheet.max_row + 1):
            if not task_sheet.cell(row=row, column=1).value:
                next_row = row
                break

        # Add the task
        task_sheet.cell(row=next_row, column=1, value=new_task_id)
        task_sheet.cell(row=next_row, column=2, value=task_text)
        task_sheet.cell(row=next_row, column=3, value=owner)
        task_sheet.cell(row=next_row, column=4, value=due_date)
        task_sheet.cell(row=next_row, column=5, value='Open')
        task_sheet.cell(row=next_row, column=6, value='Dashboard')

        wb.save(file_path)
        wb.close()

        new_task = {
            'Task ID': new_task_id,
            'Task': task_text,
            'Owner': owner,
            'Due Date': due_date,
            'Status': 'Open',
            'Source': 'Dashboard'
        }

        # Optionally add to Outlook
        outlook_result = None
        if add_to_outlook:
            outlook_result = create_outlook_task(task_text, due_date, project)

        return jsonify({
            "success": True,
            "task": new_task,
            "outlook": outlook_result
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/calendar/event', methods=['POST'])
def api_create_calendar_event():
    """Create a calendar event/meeting invite."""
    try:
        data = request.get_json()
        title = data.get('title')
        description = data.get('description', '')
        date = data.get('date')
        time = data.get('time', '09:00')
        duration = data.get('duration', 30)  # minutes
        project = data.get('project', '')
        attendee = data.get('attendee')

        if not title:
            return jsonify({"success": False, "error": "Title required"}), 400
        if not date:
            return jsonify({"success": False, "error": "Date required"}), 400

        # Try Microsoft Graph API first, fall back to ICS file
        result = create_outlook_calendar_event(title, date, time, duration, project, description, attendee)
        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/calendar/download/<filename>', methods=['GET'])
def api_download_ics(filename):
    """Download an ICS file."""
    from flask import send_file
    ics_folder = Path(__file__).parent / 'ics_files'
    file_path = ics_folder / filename

    if not file_path.exists():
        return jsonify({"success": False, "error": "File not found"}), 404

    return send_file(file_path, as_attachment=True, download_name=filename)


def create_outlook_calendar_event(title, date, time='09:00', duration=30, project='', description='', attendee=None):
    """Create a calendar event in Outlook via Microsoft Graph API."""
    import requests
    from datetime import datetime, timedelta

    # Microsoft Graph API settings from environment
    client_id = os.getenv('MS_CLIENT_ID')
    client_secret = os.getenv('MS_CLIENT_SECRET')
    tenant_id = os.getenv('MS_TENANT_ID')
    user_email = os.getenv('MS_USER_EMAIL')

    event_title = f"[{project}] {title}" if project else title

    # Parse times
    start_dt = datetime.strptime(f"{date} {time}", '%Y-%m-%d %H:%M')
    end_dt = start_dt + timedelta(minutes=duration)

    # If Graph API credentials not configured, fall back to ICS
    if not all([client_id, client_secret, tenant_id, user_email]):
        print("[Calendar] MS Graph not configured, falling back to ICS file")
        return create_calendar_ics(title, date, time, duration, project, description, attendee)

    try:
        # Get access token
        token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
        token_data = {
            'grant_type': 'client_credentials',
            'client_id': client_id,
            'client_secret': client_secret,
            'scope': 'https://graph.microsoft.com/.default'
        }
        token_response = requests.post(token_url, data=token_data)
        token_response.raise_for_status()
        access_token = token_response.json()['access_token']

        # Create event via Graph API
        event_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/calendar/events"
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json'
        }

        event_body = {
            'subject': event_title,
            'body': {
                'contentType': 'text',
                'content': description
            },
            'start': {
                'dateTime': start_dt.isoformat(),
                'timeZone': 'America/Los_Angeles'
            },
            'end': {
                'dateTime': end_dt.isoformat(),
                'timeZone': 'America/Los_Angeles'
            },
            'isReminderOn': True,
            'reminderMinutesBeforeStart': 15
        }

        # Add attendee if provided
        if attendee:
            event_body['attendees'] = [{
                'emailAddress': {
                    'address': attendee
                },
                'type': 'required'
            }]
            # Send invite to attendees
            event_body['isOnlineMeeting'] = False

        response = requests.post(event_url, headers=headers, json=event_body)
        response.raise_for_status()

        result = response.json()
        print(f"[Calendar] Event created in Outlook: {result.get('id')}")

        return {
            "success": True,
            "message": "Event created in Outlook calendar",
            "method": "graph_api",
            "event_id": result.get('id'),
            "web_link": result.get('webLink'),
            "event": {
                "title": event_title,
                "date": date,
                "time": time,
                "duration": duration
            }
        }

    except requests.exceptions.HTTPError as e:
        error_detail = e.response.json() if e.response else str(e)
        print(f"[Calendar] Graph API error: {error_detail}")
        # Fall back to ICS file
        return create_calendar_ics(title, date, time, duration, project, description, attendee)
    except Exception as e:
        print(f"[Calendar] Error: {e}, falling back to ICS")
        return create_calendar_ics(title, date, time, duration, project, description, attendee)


def create_calendar_ics(title, date, time='09:00', duration=30, project='', description='', attendee=None):
    """Create an ICS calendar invite file (fallback when Graph API not available)."""
    from datetime import datetime, timedelta
    import uuid

    uid = str(uuid.uuid4())
    now = datetime.now().strftime('%Y%m%dT%H%M%SZ')

    # Parse start time
    start_dt = datetime.strptime(f"{date} {time}", '%Y-%m-%d %H:%M')
    end_dt = start_dt + timedelta(minutes=duration)

    # Format for ICS (local time)
    start_str = start_dt.strftime('%Y%m%dT%H%M%S')
    end_str = end_dt.strftime('%Y%m%dT%H%M%S')

    event_title = f"[{project}] {title}" if project else title

    # Build ICS content
    ics_lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Risk Management System//Calendar//EN",
        "METHOD:REQUEST",
        "BEGIN:VEVENT",
        f"UID:{uid}",
        f"DTSTAMP:{now}",
        f"DTSTART:{start_str}",
        f"DTEND:{end_str}",
        f"SUMMARY:{event_title}",
        f"DESCRIPTION:{description}",
        "STATUS:CONFIRMED",
    ]

    if attendee:
        ics_lines.append(f"ATTENDEE;RSVP=TRUE:mailto:{attendee}")

    ics_lines.extend([
        "END:VEVENT",
        "END:VCALENDAR"
    ])

    ics_content = '\r\n'.join(ics_lines)

    # Save to ics_files folder
    ics_folder = Path(__file__).parent / 'ics_files'
    ics_folder.mkdir(exist_ok=True)

    filename = f'event_{uid[:8]}.ics'
    file_path = ics_folder / filename

    with open(file_path, 'w') as f:
        f.write(ics_content)

    return {
        "success": True,
        "message": "Calendar invite created (download ICS file)",
        "method": "ics_file",
        "filename": filename,
        "file_path": str(file_path),
        "event": {
            "title": event_title,
            "date": date,
            "time": time,
            "duration": duration
        }
    }


# ============== Email Processing Endpoints ==============

@app.route('/api/email/check', methods=['GET'])
def api_check_emails():
    """Check for emails with attachments (preview without processing)."""
    try:
        from email_reader import fetch_emails_with_attachments

        days_back = request.args.get('days', 7, type=int)
        subject_filter = request.args.get('subject')

        emails = fetch_emails_with_attachments(
            days_back=days_back,
            subject_filter=subject_filter
        )

        # Return summary without full content
        summary = []
        for e in emails:
            summary.append({
                'subject': e['subject'],
                'sender': e['sender'],
                'date': e['date'],
                'attachments': [
                    {'filename': a['filename'], 'type': a['type'], 'size': a['size']}
                    for a in e['attachments']
                ]
            })

        return jsonify({
            "success": True,
            "emails": summary,
            "count": len(summary)
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/email/process', methods=['POST'])
def api_process_emails():
    """Fetch emails with attachments and process them through risk extraction."""
    try:
        from email_reader import process_email_attachments

        data = request.get_json() or {}
        project = data.get('project', 'HB')
        days_back = data.get('days', 7)
        subject_filter = data.get('subject_filter')

        result = process_email_attachments(
            project_code=project,
            days_back=days_back,
            subject_filter=subject_filter
        )

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/updates', methods=['GET'])
def api_get_updates():
    """Get recent updates from the Update Log."""
    project = request.args.get('project', 'HB')
    limit = request.args.get('limit', 10, type=int)
    try:
        data = read_excel_data(project)
        # Sort by timestamp descending and limit
        updates = sorted(data['updates'], key=lambda x: x.get('Timestamp', ''), reverse=True)[:limit]
        return jsonify({
            "success": True,
            "updates": updates,
            "count": len(updates)
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/milestones', methods=['GET'])
def api_get_milestones():
    """Get P6 baseline milestones for a project."""
    project = request.args.get('project')
    try:
        if project:
            data = read_excel_data(project)
            milestones = data.get('milestones', [])
            for m in milestones:
                m['project'] = project
        else:
            # Get all milestones from all projects
            milestones = []
            for proj in get_all_projects():
                data = read_excel_data(proj)
                for m in data.get('milestones', []):
                    m['project'] = proj
                    milestones.append(m)

        # Calculate summary stats
        critical_count = len([m for m in milestones if m.get('Status') == 'Critical'])
        at_risk_count = len([m for m in milestones if m.get('Status') == 'At Risk'])
        on_track_count = len([m for m in milestones if m.get('Status') == 'On Track'])
        complete_count = len([m for m in milestones if m.get('Status') == 'Complete'])

        return jsonify({
            "success": True,
            "milestones": milestones,
            "count": len(milestones),
            "summary": {
                "critical": critical_count,
                "at_risk": at_risk_count,
                "on_track": on_track_count,
                "complete": complete_count
            }
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


def get_risk_register_path(project_code):
    """Get the path to a project's Risk Register file."""
    base_path = Path(__file__).parent.resolve()
    return base_path / 'Risk_Registers' / f'Risk_Register_{project_code}.xlsx'


def generate_report_html(project, report_type):
    """Generate HTML content for a report."""
    from datetime import datetime

    # Get project data
    if project and project != 'all':
        data = read_excel_data(project)
        projects_data = {project: data}
    else:
        # Portfolio-wide report
        projects_data = {}
        for proj in get_all_projects():
            projects_data[proj] = read_excel_data(proj)

    # Compile stats
    all_risks = []
    all_tasks = []
    all_milestones = []

    for proj, data in projects_data.items():
        for r in data['risks']:
            r['project'] = proj
            all_risks.append(r)
        for t in data['tasks']:
            t['project'] = proj
            all_tasks.append(t)
        for m in data.get('milestones', []):
            m['project'] = proj
            all_milestones.append(m)

    # Calculate stats
    active_risks = [r for r in all_risks if r.get('Status') in ['Open', 'Active', 'Escalated']]
    high_risks = [r for r in active_risks if r.get('Probability') == 'High' and r.get('Impact') == 'High']
    open_tasks = [t for t in all_tasks if t.get('Status') not in ['Completed', 'Done']]
    critical_milestones = [m for m in all_milestones if m.get('Status') in ['Critical', 'At Risk']]

    # Build HTML
    title_map = {
        'current': 'Current Status Report',
        'daily': 'Daily Digest',
        'weekly': 'Weekly Summary',
        'monthly': 'Monthly Executive Report'
    }

    title = title_map.get(report_type, 'Status Report')
    now = datetime.now()

    html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; color: #333; }}
        h1 {{ color: #1e3a5f; border-bottom: 2px solid #1e3a5f; padding-bottom: 10px; }}
        h2 {{ color: #2c5282; margin-top: 30px; }}
        .header {{ background: #1e3a5f; color: white; padding: 20px; margin: -40px -40px 30px -40px; }}
        .header h1 {{ color: white; border: none; margin: 0; }}
        .header p {{ margin: 5px 0 0 0; opacity: 0.9; }}
        .summary-box {{ background: #f7fafc; border: 1px solid #e2e8f0; padding: 15px; margin: 20px 0; border-radius: 8px; }}
        .stat {{ display: inline-block; margin-right: 40px; }}
        .stat-value {{ font-size: 28px; font-weight: bold; color: #1e3a5f; }}
        .stat-label {{ font-size: 14px; color: #666; }}
        .critical {{ color: #c53030; }}
        .warning {{ color: #d69e2e; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th, td {{ padding: 10px; text-align: left; border-bottom: 1px solid #e2e8f0; }}
        th {{ background: #edf2f7; font-weight: 600; }}
        .status-critical {{ background: #fed7d7; color: #c53030; padding: 3px 8px; border-radius: 4px; }}
        .status-at-risk {{ background: #feebc8; color: #c05621; padding: 3px 8px; border-radius: 4px; }}
        .status-open {{ background: #bee3f8; color: #2b6cb0; padding: 3px 8px; border-radius: 4px; }}
        .alert-banner {{ background: #fed7d7; border-left: 4px solid #c53030; padding: 15px; margin: 20px 0; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
        <p>Generated: {now.strftime('%B %d, %Y at %I:%M %p')}</p>
    </div>
'''

    # Critical alerts section
    if high_risks or critical_milestones:
        html += '<div class="alert-banner"><strong>CRITICAL ALERTS</strong><ul>'
        for r in high_risks[:5]:
            html += f'<li><strong>[{r.get("project")}]</strong> {r.get("Risk ID")}: {r.get("Description", "")[:100]}</li>'
        for m in critical_milestones[:5]:
            html += f'<li><strong>[{m.get("project")}]</strong> Milestone: {m.get("Milestone", "")} - {m.get("Status", "")}</li>'
        html += '</ul></div>'

    # Summary stats
    html += f'''
    <div class="summary-box">
        <div class="stat">
            <div class="stat-value">{len(active_risks)}</div>
            <div class="stat-label">Active Risks</div>
        </div>
        <div class="stat">
            <div class="stat-value critical">{len(high_risks)}</div>
            <div class="stat-label">High/High Risks</div>
        </div>
        <div class="stat">
            <div class="stat-value">{len(open_tasks)}</div>
            <div class="stat-label">Open Tasks</div>
        </div>
        <div class="stat">
            <div class="stat-value warning">{len(critical_milestones)}</div>
            <div class="stat-label">Critical Milestones</div>
        </div>
    </div>
'''

    # Project breakdown for portfolio reports
    if len(projects_data) > 1 or report_type in ['daily', 'monthly']:
        html += '<h2>Project Status</h2><table><tr><th>Project</th><th>Active Risks</th><th>High Priority</th><th>Open Tasks</th><th>Health</th></tr>'
        for proj in sorted(projects_data.keys()):
            stats = get_project_stats(proj)
            health_class = 'status-critical' if stats['health'] == 'Critical' else 'status-at-risk' if stats['health'] == 'At Risk' else 'status-open'
            html += f'<tr><td><strong>{proj}</strong></td><td>{stats["active_risks"]}</td><td>{stats["high_priority"]}</td><td>{stats["open_tasks"]}</td><td><span class="{health_class}">{stats["health"]}</span></td></tr>'
        html += '</table>'

    # Active risks table
    html += '<h2>Active Risks</h2><table><tr><th>Project</th><th>Risk ID</th><th>Description</th><th>Prob/Impact</th><th>Status</th></tr>'
    for r in sorted(active_risks, key=lambda x: (x.get('Probability') != 'High', x.get('project', '')))[:15]:
        status_class = 'status-critical' if r.get('Probability') == 'High' else 'status-open'
        html += f'<tr><td>{r.get("project", "")}</td><td>{r.get("Risk ID", "")}</td><td>{str(r.get("Description", ""))[:80]}</td><td>{r.get("Probability", "")}/{r.get("Impact", "")}</td><td><span class="{status_class}">{r.get("Status", "")}</span></td></tr>'
    html += '</table>'

    # Open tasks table
    html += '<h2>Open Tasks</h2><table><tr><th>Project</th><th>Task ID</th><th>Task</th><th>Owner</th><th>Due Date</th></tr>'
    for t in sorted(open_tasks, key=lambda x: x.get('Due Date', '') or '')[:15]:
        html += f'<tr><td>{t.get("project", "")}</td><td>{t.get("Task ID", "")}</td><td>{str(t.get("Task", ""))[:80]}</td><td>{t.get("Owner", "")}</td><td>{t.get("Due Date", "")}</td></tr>'
    html += '</table>'

    # Critical milestones
    if critical_milestones:
        html += '<h2>Critical Milestones</h2><table><tr><th>Project</th><th>Milestone</th><th>Baseline</th><th>Current</th><th>Status</th></tr>'
        for m in critical_milestones[:10]:
            status_class = 'status-critical' if m.get('Status') == 'Critical' else 'status-at-risk'
            html += f'<tr><td>{m.get("project", "")}</td><td>{m.get("Milestone", "")}</td><td>{m.get("Baseline Date", "")}</td><td>{m.get("Current Date", "")}</td><td><span class="{status_class}">{m.get("Status", "")}</span></td></tr>'
        html += '</table>'

    html += '</body></html>'
    return html


def generate_report_pdf(project, report_type):
    """Generate a PDF report using ReportLab."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from datetime import datetime
    import io

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, textColor=colors.HexColor('#1e3a5f'))
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#2c5282'))

    elements = []

    # Get data
    if project and project != 'all':
        data = read_excel_data(project)
        projects_data = {project: data}
    else:
        projects_data = {}
        for proj in get_all_projects():
            projects_data[proj] = read_excel_data(proj)

    # Compile stats
    all_risks = []
    all_tasks = []

    for proj, data in projects_data.items():
        for r in data['risks']:
            r['project'] = proj
            all_risks.append(r)
        for t in data['tasks']:
            t['project'] = proj
            all_tasks.append(t)

    active_risks = [r for r in all_risks if r.get('Status') in ['Open', 'Active', 'Escalated']]
    high_risks = [r for r in active_risks if r.get('Probability') == 'High' and r.get('Impact') == 'High']
    open_tasks = [t for t in all_tasks if t.get('Status') not in ['Completed', 'Done']]

    # Title
    title_map = {
        'current': 'Current Status Report',
        'daily': 'Daily Digest',
        'weekly': 'Weekly Summary',
        'monthly': 'Monthly Executive Report'
    }
    title = title_map.get(report_type, 'Status Report')
    now = datetime.now()

    elements.append(Paragraph(title, title_style))
    elements.append(Paragraph(f"Generated: {now.strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
    elements.append(Spacer(1, 0.3*inch))

    # Summary
    summary_data = [
        ['Active Risks', 'High/High', 'Open Tasks', 'Projects'],
        [str(len(active_risks)), str(len(high_risks)), str(len(open_tasks)), str(len(projects_data))]
    ]
    summary_table = Table(summary_data, colWidths=[1.5*inch]*4)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 1), (-1, 1), 16),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 0.3*inch))

    # Active risks
    elements.append(Paragraph("Active Risks", heading_style))
    if active_risks:
        risk_data = [['Project', 'Risk ID', 'Description', 'Prob/Impact']]
        for r in sorted(active_risks, key=lambda x: (x.get('Probability') != 'High', x.get('project', '')))[:10]:
            desc = str(r.get('Description', ''))[:50] + ('...' if len(str(r.get('Description', ''))) > 50 else '')
            risk_data.append([r.get('project', ''), r.get('Risk ID', ''), desc, f"{r.get('Probability', '')}/{r.get('Impact', '')}"])
        risk_table = Table(risk_data, colWidths=[0.8*inch, 1*inch, 3.5*inch, 1*inch])
        risk_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#edf2f7')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(risk_table)
    elements.append(Spacer(1, 0.2*inch))

    # Open tasks
    elements.append(Paragraph("Open Tasks", heading_style))
    if open_tasks:
        task_data = [['Project', 'Task ID', 'Task', 'Due Date']]
        for t in sorted(open_tasks, key=lambda x: x.get('Due Date', '') or '')[:10]:
            task_desc = str(t.get('Task', ''))[:50] + ('...' if len(str(t.get('Task', ''))) > 50 else '')
            task_data.append([t.get('project', ''), t.get('Task ID', ''), task_desc, str(t.get('Due Date', '') or '')])
        task_table = Table(task_data, colWidths=[0.8*inch, 1*inch, 3.5*inch, 1*inch])
        task_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#edf2f7')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(task_table)

    doc.build(elements)
    buffer.seek(0)
    return buffer


@app.route('/api/reports/pdf', methods=['GET'])
def api_download_report_pdf():
    """Generate and download a PDF report."""
    from flask import send_file
    from datetime import datetime

    project = request.args.get('project', 'all')
    report_type = request.args.get('type', 'current')

    try:
        pdf_buffer = generate_report_pdf(project, report_type)

        filename = f"report_{project}_{report_type}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"

        return send_file(
            pdf_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/reports/preview', methods=['GET'])
def api_preview_report():
    """Preview a report as HTML."""
    project = request.args.get('project', 'all')
    report_type = request.args.get('type', 'current')

    try:
        html = generate_report_html(project, report_type)
        return html, 200, {'Content-Type': 'text/html'}
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/reports/send', methods=['POST'])
def api_send_report():
    """Send a report via email."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders
    from datetime import datetime

    try:
        data = request.get_json()
        project = data.get('project', 'all')
        report_type = data.get('type', 'current')
        recipient_email = data.get('email')
        include_attachments = data.get('include_attachments', False)

        if not recipient_email:
            return jsonify({"success": False, "error": "Email address required"}), 400

        # Email config
        smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
        smtp_port = int(os.getenv('SMTP_PORT', 587))
        sender_email = os.getenv('EMAIL_FROM') or os.getenv('SENDER_EMAIL')
        email_password = os.getenv('EMAIL_PASSWORD')

        if not sender_email or not email_password:
            return jsonify({"success": False, "error": "Email credentials not configured"}), 500

        # Build subject and HTML
        title_map = {
            'current': 'Current Status Report',
            'daily': 'Daily Digest',
            'weekly': 'Weekly Summary',
            'monthly': 'Monthly Executive Report'
        }
        subject = f"[Risk Management] {title_map.get(report_type, 'Report')}"
        if project and project != 'all':
            subject += f" - {project}"

        html_content = generate_report_html(project, report_type)

        # Create message
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Subject'] = subject

        msg.attach(MIMEText(html_content, 'html'))

        # Attach PDF
        pdf_buffer = generate_report_pdf(project, report_type)
        pdf_attachment = MIMEBase('application', 'pdf')
        pdf_attachment.set_payload(pdf_buffer.read())
        encoders.encode_base64(pdf_attachment)
        pdf_filename = f"report_{project}_{report_type}_{datetime.now().strftime('%Y%m%d')}.pdf"
        pdf_attachment.add_header('Content-Disposition', 'attachment', filename=pdf_filename)
        msg.attach(pdf_attachment)

        # Attach Risk Registers if requested
        if include_attachments:
            if project and project != 'all':
                projects_to_attach = [project]
            else:
                projects_to_attach = get_all_projects()

            for proj in projects_to_attach:
                file_path = get_risk_register_path(proj)
                if file_path.exists():
                    with open(file_path, 'rb') as f:
                        attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                        attachment.set_payload(f.read())
                        encoders.encode_base64(attachment)
                        attachment.add_header('Content-Disposition', 'attachment', filename=f'Risk_Register_{proj}.xlsx')
                        msg.attach(attachment)

        # Send email
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, email_password)
            server.send_message(msg)

        return jsonify({
            "success": True,
            "message": f"Report sent to {recipient_email}",
            "report_type": report_type,
            "project": project,
            "attachments_included": include_attachments
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/stats', methods=['GET'])
def api_get_stats():
    """Get dashboard statistics."""
    project = request.args.get('project', 'HB')
    try:
        from datetime import datetime, timedelta
        data = read_excel_data(project)
        risks = data['risks']
        tasks = data['tasks']

        today = datetime.now().date()

        # Count risks by status
        active_risks = len([r for r in risks if r.get('Status') in ['Open', 'Active', 'Escalated']])
        watching_risks = len([r for r in risks if r.get('Status') == 'Watching'])
        closed_risks = len([r for r in risks if r.get('Status') == 'Closed'])

        # Count tasks
        open_tasks = [t for t in tasks if t.get('Status') in ['Open', 'In Progress', None, '']]

        # Overdue tasks
        overdue_count = 0
        for t in open_tasks:
            due_date = t.get('Due Date')
            if due_date:
                try:
                    if isinstance(due_date, str):
                        due_date = datetime.strptime(due_date, '%Y-%m-%d').date()
                    elif hasattr(due_date, 'date'):
                        due_date = due_date.date()
                    if due_date < today:
                        overdue_count += 1
                except:
                    pass

        # Items not green (non-closed risks)
        items_not_green = [r for r in risks if r.get('Status') not in ['Closed', None, '']]

        # High priority items
        high_priority = [r for r in risks if r.get('Probability') == 'High' and r.get('Status') != 'Closed']

        return jsonify({
            "success": True,
            "stats": {
                "active_risks": active_risks,
                "watching_risks": watching_risks,
                "closed_risks": closed_risks,
                "total_risks": len(risks),
                "open_tasks": len(open_tasks),
                "overdue_tasks": overdue_count,
                "total_tasks": len(tasks),
                "items_not_green": len(items_not_green),
                "high_priority": len(high_priority)
            },
            "last_updated": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ============== Monthly Report Endpoints ==============

@app.route('/api/monthly-report/<project>/generate', methods=['POST'])
def api_generate_monthly_report(project):
    """
    Generate and optionally send a monthly report.

    POST /api/monthly-report/RH/generate
    Body:
    {
        "email": "recipient@example.com",  // Optional - if provided, sends email
        "include_attachments": true,       // Include Word doc and Excel exports
        "preview_only": false              // If true, just returns HTML preview
    }
    """
    from monthly_report import (
        send_monthly_report,
        preview_monthly_report,
        generate_word_document,
        generate_task_export,
        read_project_data,
        calculate_project_health
    )

    try:
        data = request.get_json() or {}
        email = data.get('email')
        include_attachments = data.get('include_attachments', True)
        preview_only = data.get('preview_only', False)

        # Check project exists
        project_path = get_risk_register_path(project)
        if not project_path.exists():
            return jsonify({"success": False, "error": f"Project {project} not found"}), 404

        if preview_only:
            # Return HTML preview
            html = preview_monthly_report(project)
            return html, 200, {'Content-Type': 'text/html'}

        if email:
            # Send email with report
            result = send_monthly_report(project, email, include_attachments=include_attachments)
            return jsonify(result), 200 if result['success'] else 500
        else:
            # Generate preview without sending
            html = preview_monthly_report(project)
            return jsonify({
                "success": True,
                "message": "Report generated (preview mode - no email sent)",
                "project": project,
                "preview_html": html[:500] + "...",
                "hint": "Add 'email' field to send the report"
            })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/monthly-report/<project>/preview', methods=['GET'])
def api_preview_monthly_report(project):
    """Preview the monthly report as HTML."""
    from monthly_report import preview_monthly_report

    try:
        project_path = get_risk_register_path(project)
        if not project_path.exists():
            return jsonify({"success": False, "error": f"Project {project} not found"}), 404

        html = preview_monthly_report(project)
        return html, 200, {'Content-Type': 'text/html'}

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/monthly-report/<project>/download/docx', methods=['GET'])
def api_download_monthly_docx(project):
    """Download monthly report as Word document."""
    from flask import send_file
    from monthly_report import generate_word_document, read_project_data, calculate_project_health
    from datetime import datetime

    try:
        project_path = get_risk_register_path(project)
        if not project_path.exists():
            return jsonify({"success": False, "error": f"Project {project} not found"}), 404

        data = read_project_data(project)
        health = calculate_project_health(data)
        report_date = datetime.now()

        doc_buffer = generate_word_document(project, data, health, report_date)

        filename = f"Monthly_Report_{project}_{report_date.strftime('%Y%m')}.docx"
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/monthly-report/<project>/download/tasks', methods=['GET'])
def api_download_task_export(project):
    """Download task list as Excel."""
    from flask import send_file
    from monthly_report import generate_task_export, read_project_data
    from datetime import datetime

    try:
        project_path = get_risk_register_path(project)
        if not project_path.exists():
            return jsonify({"success": False, "error": f"Project {project} not found"}), 404

        data = read_project_data(project)
        excel_buffer = generate_task_export(project, data)

        filename = f"Task_List_{project}_{datetime.now().strftime('%Y%m')}.xlsx"
        return send_file(
            excel_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


def scheduled_daily_digest():
    """Scheduled job to send daily digest at 6:00 AM."""
    print("[Scheduler] Running daily digest job...", flush=True)
    try:
        recipients = os.getenv('EMAIL_TO', '')
        projects = get_all_projects()
        for project in projects:
            send_digest_email(project, recipients=recipients)
    except Exception as e:
        print(f"[Scheduler] Error sending daily digest: {e}", flush=True)


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', 'false').lower() == 'true'

    print(f"Starting Risk Management Server on port {port}")
    print(f"Debug mode: {debug}")
    print("\nEndpoints:")
    print(f"  GET  /health         - Health check")
    print(f"  GET  /projects       - List available projects")
    print(f"  POST /process        - Process content and extract risks/tasks")
    print(f"  POST /digest         - Send digest email now")
    print(f"  GET  /digest/preview - Preview digest in browser")
    print("\nExample request:")
    print('''  curl -X POST http://localhost:5000/process \\
    -H "Content-Type: application/json" \\
    -d '{"project": "YOUR_PROJECT", "content": "...", "source_type": "meeting", "source_name": "Weekly Standup"}'
''')

    # Start folder watchers for all projects
    base_path = Path(__file__).parent.resolve()
    observers = []
    for project in get_all_projects():
        transcripts_path = base_path / project / "Transcripts"
        if transcripts_path.exists():
            observer = start_folder_watcher(transcripts_path, project)
            observers.append(observer)

    # Start scheduler for daily digest at 6:00 AM
    scheduler = BackgroundScheduler()
    scheduler.add_job(
        func=scheduled_daily_digest,
        trigger='cron',
        hour=6,
        minute=0,
        id='daily_digest',
        name='Send daily risk digest email',
        replace_existing=True
    )
    scheduler.start()
    print("[Scheduler] Daily digest scheduled for 6:00 AM")

    # Ensure scheduler shuts down on exit
    atexit.register(lambda: scheduler.shutdown())

    try:
        app.run(host='0.0.0.0', port=port, debug=debug)
    finally:
        observer.stop()
        observer.join()
        scheduler.shutdown()
