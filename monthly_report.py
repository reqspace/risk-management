"""
Monthly Report Generator for MFI LDES Projects
Generates HTML emails, Word documents, and Excel task exports.
Uses MFI branding: Cyan #33A9DC, Dark Gray #58595B
"""

import os
import base64
import smtplib
from pathlib import Path
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from io import BytesIO

import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

# MFI Brand Colors
MFI_CYAN = '#33A9DC'
MFI_DARK_GRAY = '#58595B'
MFI_CYAN_RGB = RGBColor(0x33, 0xA9, 0xDC)
MFI_GRAY_RGB = RGBColor(0x58, 0x59, 0x5B)

# Project name mapping
PROJECT_NAMES = {
    'RH': 'Rams Hill - 4MWh LDES Microgrid',
    'MCBCP': 'MCB Camp Pendleton - 400MWh Military Installation',
    'CEC': 'California Energy Commission - LDES Program',
    'HB': 'Hummingbird Project'
}

BASE_PATH = Path(__file__).parent.resolve()
LOGO_PATH = BASE_PATH / 'MFI_Logo.png'


def get_logo_base64():
    """Read the MFI logo and return as base64 string."""
    if LOGO_PATH.exists():
        with open(LOGO_PATH, 'rb') as f:
            return base64.b64encode(f.read()).decode('utf-8')
    return None


def read_project_data(project_code):
    """Read all data from a project's Risk Register."""
    file_path = BASE_PATH / 'Risk_Registers' / f'Risk_Register_{project_code}.xlsx'

    if not file_path.exists():
        return {'risks': [], 'tasks': [], 'milestones': [], 'updates': []}

    wb = openpyxl.load_workbook(file_path)

    def parse_date(val):
        if val is None:
            return None
        if isinstance(val, datetime):
            return val.strftime('%Y-%m-%d')
        return str(val)

    # Read risks
    risks = []
    risk_sheet = wb['Risk Register']
    headers = [cell.value for cell in risk_sheet[1]]

    for row_num in range(3, risk_sheet.max_row + 1):
        row_data = {}
        for i, cell in enumerate(risk_sheet[row_num]):
            if i < len(headers) and headers[i]:
                val = cell.value
                if headers[i] in ['Date Identified', 'Last Updated', 'Closed Date']:
                    val = parse_date(val)
                row_data[headers[i]] = val
        if row_data.get('Risk ID'):
            risks.append(row_data)

    # Read tasks
    tasks = []
    task_sheet = wb['Tasks']
    task_headers = [cell.value for cell in task_sheet[1]]

    for row_num in range(3, task_sheet.max_row + 1):
        row_data = {}
        for i, cell in enumerate(task_sheet[row_num]):
            if i < len(task_headers) and task_headers[i]:
                val = cell.value
                if task_headers[i] in ['Due Date', 'Created Date', 'Completed Date']:
                    val = parse_date(val)
                row_data[task_headers[i]] = val
        if row_data.get('Task ID'):
            tasks.append(row_data)

    # Read milestones
    milestones = []
    if 'Milestones' in wb.sheetnames:
        ms_sheet = wb['Milestones']
        ms_headers = [cell.value for cell in ms_sheet[1]]

        for row_num in range(2, ms_sheet.max_row + 1):
            row_data = {}
            for i, cell in enumerate(ms_sheet[row_num]):
                if i < len(ms_headers) and ms_headers[i]:
                    val = cell.value
                    if ms_headers[i] in ['Baseline Date', 'Current Date']:
                        val = parse_date(val)
                    row_data[ms_headers[i]] = val
            if row_data.get('Milestone ID') or row_data.get('Milestone'):
                milestones.append(row_data)

    # Read updates
    updates = []
    if 'Update Log' in wb.sheetnames:
        update_sheet = wb['Update Log']
        update_headers = [cell.value for cell in update_sheet[1]]

        for row_num in range(3, update_sheet.max_row + 1):
            row_data = {}
            for i, cell in enumerate(update_sheet[row_num]):
                if i < len(update_headers) and update_headers[i]:
                    row_data[update_headers[i]] = cell.value
            if row_data.get('Timestamp'):
                updates.append(row_data)

    wb.close()
    return {'risks': risks, 'tasks': tasks, 'milestones': milestones, 'updates': updates}


def calculate_project_health(data):
    """Calculate project health status based on risks and milestones."""
    risks = data['risks']
    milestones = data['milestones']
    tasks = data['tasks']

    active_risks = [r for r in risks if r.get('Status') in ['Open', 'Active', 'Escalated']]
    high_high_risks = [r for r in active_risks if r.get('Probability') == 'High' and r.get('Impact') == 'High']
    critical_milestones = [m for m in milestones if m.get('Status') in ['Critical', 'At Risk']]
    open_tasks = [t for t in tasks if t.get('Status') not in ['Completed', 'Done']]

    # Determine overdue tasks
    today = datetime.now().date()
    overdue_tasks = []
    for t in open_tasks:
        due_date = t.get('Due Date')
        if due_date:
            try:
                if isinstance(due_date, str):
                    due_date = datetime.strptime(due_date, '%Y-%m-%d').date()
                elif hasattr(due_date, 'date'):
                    due_date = due_date.date()
                if due_date < today:
                    overdue_tasks.append(t)
            except:
                pass

    # Health calculation
    if len(high_high_risks) >= 2 or len(critical_milestones) >= 2:
        health = 'Critical'
        health_color = '#c53030'
    elif len(high_high_risks) >= 1 or len(critical_milestones) >= 1 or len(overdue_tasks) >= 3:
        health = 'At Risk'
        health_color = '#d69e2e'
    elif len(active_risks) >= 3 or len(overdue_tasks) >= 1:
        health = 'Caution'
        health_color = '#3182ce'
    else:
        health = 'On Track'
        health_color = '#38a169'

    return {
        'status': health,
        'color': health_color,
        'active_risks': len(active_risks),
        'high_high_risks': len(high_high_risks),
        'critical_milestones': len(critical_milestones),
        'open_tasks': len(open_tasks),
        'overdue_tasks': len(overdue_tasks),
        'completed_tasks': len([t for t in tasks if t.get('Status') in ['Completed', 'Done']])
    }


def generate_executive_summary(project_code, data, health):
    """Generate executive summary text based on project data."""
    project_name = PROJECT_NAMES.get(project_code, project_code)

    # Find critical items
    risks = data['risks']
    high_risks = [r for r in risks if r.get('Probability') == 'High' and r.get('Status') != 'Closed']

    summary_lines = []

    if health['status'] == 'Critical':
        summary_lines.append(f"CRITICAL: {project_name} requires immediate attention.")
    elif health['status'] == 'At Risk':
        summary_lines.append(f"ATTENTION: {project_name} has items requiring management focus.")
    else:
        summary_lines.append(f"{project_name} is progressing as planned.")

    summary_lines.append(f"\nProject has {health['active_risks']} active risks, {health['open_tasks']} open tasks, and {health['completed_tasks']} completed tasks this period.")

    if high_risks:
        summary_lines.append(f"\nTop Risk: {high_risks[0].get('Description', 'N/A')[:150]}")
        if high_risks[0].get('Mitigation'):
            summary_lines.append(f"Mitigation: {high_risks[0].get('Mitigation', '')[:150]}")

    # Add milestone summary
    milestones = data['milestones']
    upcoming = [m for m in milestones if m.get('Status') in ['On Track', 'At Risk']]
    if upcoming:
        summary_lines.append(f"\nNext Key Milestone: {upcoming[0].get('Milestone', 'N/A')}")

    return '\n'.join(summary_lines)


def generate_html_email(project_code, data, health, report_date=None):
    """Generate HTML email with embedded MFI logo - matches dashboard grey theme."""
    if report_date is None:
        report_date = datetime.now()

    project_name = PROJECT_NAMES.get(project_code, project_code)
    logo_b64 = get_logo_base64()

    # Extract data
    risks = data['risks']
    tasks = data['tasks']
    milestones = data['milestones']

    active_risks = [r for r in risks if r.get('Status') in ['Open', 'Active', 'Escalated']]
    high_risks = [r for r in active_risks if r.get('Probability') == 'High']
    open_tasks = [t for t in tasks if t.get('Status') not in ['Completed', 'Done']]
    critical_milestones = [m for m in milestones if m.get('Status') in ['Critical', 'At Risk']]

    executive_summary = generate_executive_summary(project_code, data, health)

    # Colors matching dashboard
    SLATE_800 = '#1e293b'
    SLATE_700 = '#334155'
    SLATE_600 = '#475569'
    RED_500 = '#ef4444'
    YELLOW_500 = '#eab308'
    BLUE_500 = '#3b82f6'
    ORANGE_500 = '#f97316'
    GREEN_500 = '#22c55e'

    html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Monthly Report - {project_code}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            margin: 0;
            padding: 0;
            background: #f1f5f9;
            color: #334155;
            line-height: 1.5;
        }}
        .container {{
            max-width: 700px;
            margin: 20px auto;
            background: white;
            border-radius: 12px;
            overflow: hidden;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .header {{
            background: linear-gradient(135deg, {SLATE_800} 0%, {SLATE_700} 100%);
            padding: 24px 32px;
            color: white;
        }}
        .header-row {{
            display: flex;
            align-items: center;
            justify-content: space-between;
        }}
        .logo {{
            height: 50px;
        }}
        .header h1 {{
            margin: 0;
            font-size: 22px;
            font-weight: 600;
        }}
        .header .subtitle {{
            margin: 4px 0 0 0;
            opacity: 0.8;
            font-size: 14px;
        }}
        .header .date {{
            font-size: 13px;
            opacity: 0.7;
            margin-top: 2px;
        }}
        .summary-bar {{
            background: {SLATE_700};
            padding: 20px 32px;
            display: flex;
            justify-content: space-between;
            flex-wrap: wrap;
            gap: 16px;
        }}
        .stat {{
            text-align: center;
            flex: 1;
            min-width: 80px;
        }}
        .stat-value {{
            font-size: 32px;
            font-weight: 700;
            line-height: 1;
        }}
        .stat-value.red {{ color: #f87171; }}
        .stat-value.yellow {{ color: #fbbf24; }}
        .stat-value.blue {{ color: #60a5fa; }}
        .stat-value.orange {{ color: #fb923c; }}
        .stat-value.green {{ color: #4ade80; }}
        .stat-value.white {{ color: white; }}
        .stat-label {{
            font-size: 11px;
            color: #94a3b8;
            margin-top: 4px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .content {{
            padding: 24px 32px;
        }}
        .section {{
            margin-bottom: 28px;
        }}
        .section-title {{
            color: {SLATE_800};
            font-size: 16px;
            font-weight: 600;
            margin-bottom: 12px;
            padding-bottom: 8px;
            border-bottom: 2px solid #e2e8f0;
            display: flex;
            align-items: center;
            gap: 8px;
        }}
        .section-title::before {{
            content: '';
            width: 4px;
            height: 18px;
            background: {MFI_CYAN};
            border-radius: 2px;
        }}
        .health-badge {{
            display: inline-block;
            padding: 6px 16px;
            border-radius: 16px;
            font-weight: 600;
            font-size: 13px;
            color: white;
            background: {health['color']};
        }}
        .executive-summary {{
            background: #f8fafc;
            padding: 16px 20px;
            border-radius: 8px;
            border-left: 4px solid {MFI_CYAN};
            white-space: pre-line;
            font-size: 14px;
            color: #475569;
        }}
        .alert-box {{
            background: #fef2f2;
            border-left: 4px solid #dc2626;
            padding: 12px 16px;
            margin: 12px 0;
            border-radius: 6px;
            font-size: 13px;
        }}
        .alert-box.warning {{
            background: #fffbeb;
            border-left-color: #d97706;
        }}
        .alert-box strong {{
            color: #991b1b;
        }}
        .alert-box.warning strong {{
            color: #92400e;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 12px 0;
            font-size: 13px;
        }}
        th {{
            background: {SLATE_800};
            color: white;
            padding: 10px 12px;
            text-align: left;
            font-weight: 500;
            font-size: 11px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        td {{
            padding: 10px 12px;
            border-bottom: 1px solid #e2e8f0;
            color: #475569;
        }}
        tr:nth-child(even) {{
            background: #f8fafc;
        }}
        .badge {{
            display: inline-block;
            padding: 4px 12px;
            border-radius: 12px;
            font-size: 11px;
            font-weight: 600;
            white-space: nowrap;
        }}
        .badge-critical {{ background: #fee2e2; color: #dc2626; }}
        .badge-at-risk {{ background: #fef3c7; color: #d97706; }}
        .badge-open {{ background: #dbeafe; color: #2563eb; }}
        .badge-on-track {{ background: #dcfce7; color: #16a34a; }}
        .badge-high {{ background: #fee2e2; color: #dc2626; }}
        .badge-medium {{ background: #fef3c7; color: #d97706; }}
        .badge-low {{ background: #dcfce7; color: #16a34a; }}
        .badge-complete {{ background: #d1fae5; color: #059669; }}
        .footer {{
            background: {SLATE_800};
            color: #94a3b8;
            padding: 16px 32px;
            text-align: center;
            font-size: 12px;
        }}
        .footer a {{
            color: {MFI_CYAN};
            text-decoration: none;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="header-row">
                <div>
                    <h1>LDES Microgrid Status Report</h1>
                    <p class="subtitle">{project_name}</p>
                    <p class="date">{report_date.strftime('%B %d, %Y')}</p>
                </div>
                {'<img src="data:image/png;base64,' + logo_b64 + '" class="logo" alt="MFI Logo">' if logo_b64 else ''}
            </div>
        </div>

        <div class="summary-bar">
            <div class="stat">
                <div class="stat-value red">{health['active_risks']}</div>
                <div class="stat-label">Active Risks</div>
            </div>
            <div class="stat">
                <div class="stat-value yellow">{health['high_high_risks']}</div>
                <div class="stat-label">High/High</div>
            </div>
            <div class="stat">
                <div class="stat-value blue">{health['open_tasks']}</div>
                <div class="stat-label">Open Tasks</div>
            </div>
            <div class="stat">
                <div class="stat-value orange">{health['overdue_tasks']}</div>
                <div class="stat-label">Overdue</div>
            </div>
            <div class="stat">
                <div class="stat-value green">{health['completed_tasks']}</div>
                <div class="stat-label">Completed</div>
            </div>
        </div>

        <div class="content">
            <!-- Health Status -->
            <div class="section">
                <div class="section-title">Project Health</div>
                <span class="health-badge">{health['status']}</span>
            </div>

            <!-- Executive Summary -->
            <div class="section">
                <div class="section-title">Executive Summary</div>
                <div class="executive-summary">{executive_summary}</div>
            </div>
'''

    # Critical Alerts
    if health['high_high_risks'] > 0 or health['critical_milestones'] > 0:
        html += '''
            <div class="section">
                <div class="section-title">Critical Alerts</div>
'''
        for r in high_risks[:3]:
            html += f'''
                <div class="alert-box">
                    <strong>{r.get('Risk ID', '')}</strong>: {str(r.get('Description', ''))[:200]}
                    <br><small>Mitigation: {str(r.get('Mitigation', 'TBD'))[:150]}</small>
                </div>
'''
        for m in critical_milestones[:2]:
            m_status = m.get('Status', '')
            m_badge = 'badge-critical' if m_status == 'Critical' else 'badge-at-risk'
            html += f'''
                <div class="alert-box warning">
                    <strong>Milestone:</strong> {m.get('Milestone', '')} - <span class="badge {m_badge}">{m_status}</span>
                    <br><small>Baseline: {m.get('Baseline Date', 'TBD')} | Current: {m.get('Current Date', 'TBD')}</small>
                </div>
'''
        html += '</div>'

    # Schedule/Milestones Section
    html += '''
            <div class="section">
                <div class="section-title">Schedule & Milestones</div>
                <table>
                    <tr>
                        <th>Milestone</th>
                        <th>Baseline</th>
                        <th>Current</th>
                        <th>Status</th>
                    </tr>
'''
    for m in milestones[:8]:
        status = m.get('Status', 'On Track')
        if status == 'Critical':
            badge_class = 'badge-critical'
        elif status == 'At Risk':
            badge_class = 'badge-at-risk'
        elif status in ['Complete', 'Completed']:
            badge_class = 'badge-complete'
        elif status == 'On Track':
            badge_class = 'badge-on-track'
        else:
            badge_class = 'badge-open'
        html += f'''
                    <tr>
                        <td>{m.get('Milestone', '')}</td>
                        <td>{m.get('Baseline Date', '')}</td>
                        <td>{m.get('Current Date', '')}</td>
                        <td><span class="badge {badge_class}">{status}</span></td>
                    </tr>
'''
    html += '''
                </table>
            </div>
'''

    # Risks Section
    html += '''
            <div class="section">
                <div class="section-title">Risk Register Summary</div>
                <table>
                    <tr>
                        <th>ID</th>
                        <th>Description</th>
                        <th>Prob/Impact</th>
                        <th>Status</th>
                    </tr>
'''
    for r in active_risks[:10]:
        prob = r.get('Probability', '')
        r_status = r.get('Status', '')
        if prob == 'High':
            prob_badge = 'badge-high'
        elif prob == 'Medium':
            prob_badge = 'badge-medium'
        else:
            prob_badge = 'badge-low'
        if r_status in ['Open', 'Active']:
            status_badge = 'badge-critical'
        elif r_status == 'Escalated':
            status_badge = 'badge-high'
        else:
            status_badge = 'badge-open'
        html += f'''
                    <tr>
                        <td>{r.get('Risk ID', '')}</td>
                        <td>{str(r.get('Description', ''))[:80]}</td>
                        <td><span class="badge {prob_badge}">{prob}/{r.get('Impact', '')}</span></td>
                        <td><span class="badge {status_badge}">{r_status}</span></td>
                    </tr>
'''
    html += '''
                </table>
            </div>
'''

    # Tasks Section
    html += '''
            <div class="section">
                <div class="section-title">Action Items</div>
                <table>
                    <tr>
                        <th>Task</th>
                        <th>Owner</th>
                        <th>Due Date</th>
                        <th>Status</th>
                    </tr>
'''
    for t in open_tasks[:10]:
        t_status = t.get('Status', 'Open')
        html += f'''
                    <tr>
                        <td>{str(t.get('Task', ''))[:60]}</td>
                        <td>{t.get('Owner', '')}</td>
                        <td>{t.get('Due Date', '')}</td>
                        <td><span class="badge badge-open">{t_status}</span></td>
                    </tr>
'''
    html += '''
                </table>
            </div>
'''

    # Safety Section (placeholder)
    html += '''
            <div class="section">
                <div class="section-title">Safety & Compliance</div>
                <p style="color: #38a169;">No safety incidents reported this period.</p>
                <p>All work continues in compliance with applicable regulations and safety protocols.</p>
            </div>
'''

    # Next Steps
    html += '''
            <div class="section">
                <div class="section-title">Next Steps</div>
                <ul>
'''
    for t in open_tasks[:5]:
        html += f'<li>{str(t.get("Task", ""))[:100]}</li>'
    html += '''
                </ul>
            </div>
        </div>

        <div class="footer">
            <p>MFI LDES Risk Management System</p>
            <p>Report generated: ''' + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '''</p>
        </div>
    </div>
</body>
</html>'''

    return html


def generate_word_document(project_code, data, health, report_date=None):
    """Generate a Word document report with MFI branding."""
    if report_date is None:
        report_date = datetime.now()

    project_name = PROJECT_NAMES.get(project_code, project_code)

    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add logo to header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    if LOGO_PATH.exists():
        run = header_para.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.5))
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Title
    title = doc.add_heading('Monthly Status Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = MFI_CYAN_RGB

    # Project info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f'{project_name}\n{report_date.strftime("%B %Y")}')
    info_run.font.size = Pt(14)
    info_run.font.color.rgb = MFI_GRAY_RGB

    doc.add_paragraph()

    # Project Health
    health_heading = doc.add_heading('Project Health', level=1)
    for run in health_heading.runs:
        run.font.color.rgb = MFI_CYAN_RGB

    health_para = doc.add_paragraph()
    health_run = health_para.add_run(f"Status: {health['status']}")
    health_run.bold = True
    health_run.font.size = Pt(14)

    # Stats table
    stats_table = doc.add_table(rows=2, cols=4)
    stats_table.style = 'Table Grid'

    headers = ['Active Risks', 'High/High Risks', 'Open Tasks', 'Completed']
    values = [str(health['active_risks']), str(health['high_high_risks']),
              str(health['open_tasks']), str(health['completed_tasks'])]

    for i, header_text in enumerate(headers):
        cell = stats_table.rows[0].cells[i]
        cell.text = header_text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for i, value in enumerate(values):
        cell = stats_table.rows[1].cells[i]
        cell.text = value
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(16)
                run.font.bold = True

    doc.add_paragraph()

    # Executive Summary
    exec_heading = doc.add_heading('Executive Summary', level=1)
    for run in exec_heading.runs:
        run.font.color.rgb = MFI_CYAN_RGB

    exec_summary = generate_executive_summary(project_code, data, health)
    doc.add_paragraph(exec_summary)

    # Schedule & Milestones
    schedule_heading = doc.add_heading('Schedule & Milestones', level=1)
    for run in schedule_heading.runs:
        run.font.color.rgb = MFI_CYAN_RGB

    milestones = data['milestones']
    if milestones:
        ms_table = doc.add_table(rows=1, cols=4)
        ms_table.style = 'Table Grid'

        header_cells = ms_table.rows[0].cells
        for i, text in enumerate(['Milestone', 'Baseline', 'Current', 'Status']):
            header_cells[i].text = text
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        for m in milestones[:10]:
            row = ms_table.add_row().cells
            row[0].text = str(m.get('Milestone', ''))[:50]
            row[1].text = str(m.get('Baseline Date', ''))
            row[2].text = str(m.get('Current Date', ''))
            row[3].text = str(m.get('Status', ''))

    doc.add_paragraph()

    # Budget Section (placeholder)
    budget_heading = doc.add_heading('Budget Status', level=1)
    for run in budget_heading.runs:
        run.font.color.rgb = MFI_CYAN_RGB
    doc.add_paragraph('Budget tracking is maintained separately. Contact project manager for detailed financial status.')

    # Risk Register
    risk_heading = doc.add_heading('Risk Register Summary', level=1)
    for run in risk_heading.runs:
        run.font.color.rgb = MFI_CYAN_RGB

    risks = data['risks']
    active_risks = [r for r in risks if r.get('Status') in ['Open', 'Active', 'Escalated']]

    if active_risks:
        risk_table = doc.add_table(rows=1, cols=4)
        risk_table.style = 'Table Grid'

        header_cells = risk_table.rows[0].cells
        for i, text in enumerate(['ID', 'Description', 'Prob/Impact', 'Mitigation']):
            header_cells[i].text = text
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        for r in active_risks[:10]:
            row = risk_table.add_row().cells
            row[0].text = str(r.get('Risk ID', ''))
            row[1].text = str(r.get('Description', ''))[:60]
            row[2].text = f"{r.get('Probability', '')}/{r.get('Impact', '')}"
            row[3].text = str(r.get('Mitigation', ''))[:60]

    doc.add_paragraph()

    # Safety Section
    safety_heading = doc.add_heading('Safety & Compliance', level=1)
    for run in safety_heading.runs:
        run.font.color.rgb = MFI_CYAN_RGB
    doc.add_paragraph('No safety incidents reported this period. All work continues in compliance with applicable regulations.')

    # Next Steps
    next_heading = doc.add_heading('Next Steps', level=1)
    for run in next_heading.runs:
        run.font.color.rgb = MFI_CYAN_RGB

    tasks = data['tasks']
    open_tasks = [t for t in tasks if t.get('Status') not in ['Completed', 'Done']]

    for t in open_tasks[:5]:
        doc.add_paragraph(str(t.get('Task', '')), style='List Bullet')

    # Attachments note
    doc.add_paragraph()
    attach_heading = doc.add_heading('Attachments', level=1)
    for run in attach_heading.runs:
        run.font.color.rgb = MFI_CYAN_RGB
    doc.add_paragraph('• Task List Export (.xlsx)')
    doc.add_paragraph('• Risk Register (.xlsx)')

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_task_export(project_code, data):
    """Generate Excel export of tasks."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Task List'

    # Header styling
    header_fill = PatternFill(start_color='33A9DC', end_color='33A9DC', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF', size=11)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Headers
    headers = ['Task ID', 'Task', 'Owner', 'Due Date', 'Status', 'Priority', 'Source']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # Data
    tasks = data['tasks']
    for row_num, task in enumerate(tasks, 2):
        ws.cell(row=row_num, column=1, value=task.get('Task ID', ''))
        ws.cell(row=row_num, column=2, value=task.get('Task', ''))
        ws.cell(row=row_num, column=3, value=task.get('Owner', ''))
        ws.cell(row=row_num, column=4, value=task.get('Due Date', ''))
        ws.cell(row=row_num, column=5, value=task.get('Status', ''))
        ws.cell(row=row_num, column=6, value=task.get('Priority', ''))
        ws.cell(row=row_num, column=7, value=task.get('Source', ''))

        for col in range(1, 8):
            ws.cell(row=row_num, column=col).border = thin_border

    # Column widths
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 50
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 10
    ws.column_dimensions['G'].width = 15

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def send_monthly_report(project_code, recipients, report_date=None, include_attachments=True):
    """
    Generate and send the monthly report via email.

    Args:
        project_code: Project code (RH, MCBCP, CEC, etc.)
        recipients: Email address or list of addresses
        report_date: Date for the report (defaults to now)
        include_attachments: Whether to include Word doc and Excel exports

    Returns:
        dict with success status and details
    """
    if report_date is None:
        report_date = datetime.now()

    if isinstance(recipients, str):
        recipients = [recipients]

    # Load project data
    data = read_project_data(project_code)
    if not data['risks'] and not data['tasks']:
        return {'success': False, 'error': f'No data found for project {project_code}'}

    health = calculate_project_health(data)
    project_name = PROJECT_NAMES.get(project_code, project_code)

    # Generate content
    html_content = generate_html_email(project_code, data, health, report_date)

    # Email configuration
    smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
    smtp_port = int(os.getenv('SMTP_PORT', 587))
    sender_email = os.getenv('EMAIL_FROM') or os.getenv('SENDER_EMAIL')
    email_password = os.getenv('EMAIL_PASSWORD')

    if not sender_email or not email_password:
        return {'success': False, 'error': 'Email credentials not configured in .env (need EMAIL_FROM and EMAIL_PASSWORD)'}

    # Build email
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = ', '.join(recipients)
    msg['Subject'] = f'[MFI] Monthly Report - {project_code} - {report_date.strftime("%B %Y")}'

    msg.attach(MIMEText(html_content, 'html'))

    attachments_added = []

    if include_attachments:
        # Add Word document
        try:
            word_buffer = generate_word_document(project_code, data, health, report_date)
            word_attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
            word_attachment.set_payload(word_buffer.read())
            encoders.encode_base64(word_attachment)
            word_filename = f'Monthly_Report_{project_code}_{report_date.strftime("%Y%m")}.docx'
            word_attachment.add_header('Content-Disposition', 'attachment', filename=word_filename)
            msg.attach(word_attachment)
            attachments_added.append(word_filename)
        except Exception as e:
            print(f"Warning: Could not generate Word document: {e}")

        # Add task export
        try:
            excel_buffer = generate_task_export(project_code, data)
            excel_attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            excel_attachment.set_payload(excel_buffer.read())
            encoders.encode_base64(excel_attachment)
            excel_filename = f'Task_List_{project_code}_{report_date.strftime("%Y%m")}.xlsx'
            excel_attachment.add_header('Content-Disposition', 'attachment', filename=excel_filename)
            msg.attach(excel_attachment)
            attachments_added.append(excel_filename)
        except Exception as e:
            print(f"Warning: Could not generate task export: {e}")

        # Add Risk Register
        risk_register_path = BASE_PATH / 'Risk_Registers' / f'Risk_Register_{project_code}.xlsx'
        if risk_register_path.exists():
            with open(risk_register_path, 'rb') as f:
                risk_attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                risk_attachment.set_payload(f.read())
                encoders.encode_base64(risk_attachment)
                risk_attachment.add_header('Content-Disposition', 'attachment', filename=f'Risk_Register_{project_code}.xlsx')
                msg.attach(risk_attachment)
                attachments_added.append(f'Risk_Register_{project_code}.xlsx')

    # Send email
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, email_password)
            server.send_message(msg)

        return {
            'success': True,
            'message': f'Monthly report sent to {", ".join(recipients)}',
            'project': project_code,
            'health_status': health['status'],
            'attachments': attachments_added,
            'report_date': report_date.strftime('%Y-%m-%d')
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}


def preview_monthly_report(project_code, report_date=None):
    """Generate HTML preview of the monthly report."""
    if report_date is None:
        report_date = datetime.now()

    data = read_project_data(project_code)
    health = calculate_project_health(data)

    return generate_html_email(project_code, data, health, report_date)


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("Usage: python monthly_report.py <project_code> [email] [--preview]")
        print("Example: python monthly_report.py RH mike@example.com")
        print("         python monthly_report.py RH --preview")
        sys.exit(1)

    project = sys.argv[1]

    if '--preview' in sys.argv:
        html = preview_monthly_report(project)
        preview_path = BASE_PATH / f'monthly_report_preview_{project}.html'
        with open(preview_path, 'w') as f:
            f.write(html)
        print(f"Preview saved to: {preview_path}")
    elif len(sys.argv) >= 3 and not sys.argv[2].startswith('--'):
        email = sys.argv[2]
        result = send_monthly_report(project, email)
        if result['success']:
            print(f"Success: {result['message']}")
            print(f"Attachments: {', '.join(result.get('attachments', []))}")
        else:
            print(f"Error: {result['error']}")
    else:
        # Just generate preview
        html = preview_monthly_report(project)
        print(html[:500] + "...")
